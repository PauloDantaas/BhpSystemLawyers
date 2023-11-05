import customtkinter as ctk
import tkinter as tk
from CTkMessagebox import CTkMessagebox
from tkinter import filedialog, scrolledtext
from PIL import Image
from docx import Document
from docx.shared import Pt
import os
import MySQLdb
from dotenv import load_dotenv
from elements import *


env_path = r'./assets/verify_identity/dont_env/env/.env'






#arquivo Principal



load_dotenv(dotenv_path=env_path)

#conexao com o banco de dados
def Connect():
    try:
        mydb = str()
        mydb = MySQLdb.connect(
        host= os.getenv("DB_HOST"),
        user=os.getenv("DB_USERNAME"),
        passwd= os.getenv("DB_PASSWORD"),
        db= os.getenv("DB_NAME"),
        autocommit = True,
        ssl_mode = "VERIFY_IDENTITY",
        ssl      = {
            "ca": "./assets/verify_identity/cacert-2023-08-22.pem"
        }
)
    
    except Exception as error:
        CTkMessagebox(title='Error', message=error, icon='cancel')

    finally:
        return mydb
    




empresas=["ANHANGUERA EDUCACIONAL PARTICIPACOES S/A", "ATIVOS S.A. SECURITIZADORA DE CREDITOS FINANCEIROS", "ATLANTICO FUNDO DE INVESTIMENTOS EM DIREITOS CREDITORIOS -  NAO PADRONIZADO", 
          "AVON COSMETICOS LTDA", "AYMORÉ CREDITO, FINANCEIRA E INVESTIMENTO S.A.", "BANCO BMG S.A.", "BANCO BRADESCARD S.A.", "BANCO BRADESCO FINANCIAMENTOS S.A.", 
          "BANCO BRADESCO S.A.", "BANCO BRADESCO S.A. (BRADESCO CARTÕES)", "BANCO CACIQUE S.A.", "BANCO CETELEM S.A.", "BANCO CITIBANK S/A", "BANCO COOPERATIVO SICOOB S.A.", 
          "BANCO CSF S.A.", "BANCO DAYCOVAL S.A.", "BANCO DIGIO S.A.", "BANCO DO BRASIL S.A.", "BANCO DO ESTADO DO RIO GRANDE DO SUL S.A.", "BANCO DO NORDESTE DO BRASIL S.A.", 
          "BANCO FIBRA S.A.", "BANCO ITAUCARD S.A.", "BANCO LOSANGO S.A. – BANCO MULTIPLO", "BANCO PAN S.A.", "BANCO SANTANDER BRASIL S.A.", "BANCO SISTEMA S.A.", 
          "BANCO SOROCRED S.A. – BANCO MULTIPLO", "BANCO TOYOTA DO BRASIL S.A.", "BANCO TRIANGULO S/A", "BANCO VOLKSWAGEN S.A.", "BANCO VOTORANTIM S.A.", "BOTICARIO PRODUTOS DE BELEZA LTDA", 
          "CAIXA ECONOMICA FEDERAL", "CCB BRASIL S/A CRÉDITO FINANCIAMENTOS E INVESTIMENTOS", "CENCONSUD BRASIL COMERCIAL S.A.", "CIELO S.A.", "CLARO S.A.", "COMPANHIA BRASILEIRA DE DISTRIBUICAO (SUPERMERCADO PÃO DE AÇUCAR)", 
          "CREDIARE S/A - CREDITO, FINANCIAMENTO E INVESTIMENTO", "CREDITATIVOS SOLUÇÕES FINANCEIRAS LTDA", "CREFISA SA CREDITO FINANCIAMENTO E INVESTIMENTOS", "DACASA FINANCEIRA S/A", 
          "DMCARD MEIOS DE PAGAMENTO LTDA", "DS CARD ADMINISTRADORA DE CARTÕES DE CREDITO LTDA", "EDITORA E DISTRIBUIDORA EDUCACIONAL S/A", "FINANCEIRA ITAU CBD S.A. – CREDITO, FINANCIAMENTO E INVESTIMENTO", 
          "FUNDO DE INVESTIMENTO EM DIREITOS CREDITORIOS IPANEMA III", "FUNDO DE INVESTIMENTO EM DIREITOS CREDITORIOS MULTSEGMENTOS NPL IPANEMA VI – NÃO PADRONIZADO", 
          "FUNDO DE INVESTIMENTO EM DIREITOS CREDITORIOS NAO PADRONIZADOS NPL II", "HAPVIDA PARTICIPACÕES E INVESTIMENTO S/A", "HIPERCARD BANCO MULTIPLO S.A.", "HOEPERS RECUPERADORA DE CREDITO S/A", 
          "IRESOLVE COMPANHIA SECURITIZADORA DE CREDITOS FINANCEIROS S.A.", "ITAPEVA XI MULTICARTEIRA FUNDO DE INVESTIMENTO EM DIREITOS CREDITORIOS NAO PADRONIZADOS", "ITAPEVA XII MULTICARTEIRA FUNDO DE INVESTIMENTOS EM DIREITOS CREDITORIOS NP", 
          "ITAU UNIBANCO S.A.", "JBCRED S/A SOCIEDADE DE CREDITO, FINANCIAMENTO E INVESTIMENTO", "KIRTON BANK S.A. – BANCO MULTIPLO", "LOJAS AMERICANAS S.A.", "LOJAS RENNER S/A", 
          "LOJAS RIACHUELO S/A", "LUIZACRED S.A. SOCIEDADE DE CREDITO, FINACIAMENTO E INVESTIMENTO - FINANCEIRAS", "MARISA LOJAS S.A.", "MGW ATIVOS – GESTÃO E ADMINISTRAÇÃO DE CRÉDITOS FINANCEIROS LTDA", 
          "MIDWAY S.A. CREDITO, FINANCIAMENTO INVESTIMENTO", "MULVI INSTITUICAO DE PAGAMENTOS S.A.", "NATURA COSMETICOS S/A", "OI MOVEL S.A.", "OMNI BANCO S.A.", "PAGSEGURO INTERNET S/A", 
          "PORTOCRED S.A. CREDITO, FINANCIAMENTO E INVESTIMENTO", "RCB INVESTIMENTOS S.A.", "REALIZE CREDITO, FINANCIAMENTO E INVESTIMENTO S.A.", "REDECARD S/A (CREDICARD)", "SASCAR - TECNOLOGIA E SEGURANCA AUTOMOTIVA S/A", 
          "SKY BRASIL SERVICOS LTDA", "SOROCRED - CREDITO, FINANCIAMENTO E INVESTIMENTO S/A", "TELEFÔNICA BRASIL S.A.", "TELEFONICA DATA S.A.", "TELEMAR NORTE LESTE S/A", "TIM S.A.", 
          "ULTRA SOM SERVIÇOS MEDICOS S/A", "VIA VAREJO S.A."]



arquivo = None

class App():
    def __init__(self):
        self.window = ctk.CTk()
        self.window.title('BHP System Lawyers')
        self.window.geometry(Elements.GEOMETRY)
        self.window.resizable(0, 0)
        self.thema_default = ctk.set_appearance_mode('light')
        self.color_default = ctk.set_default_color_theme('blue')
        self.Login()
        
        self.img_Home = ctk.CTkImage(light_image=Image.open('./assets/icons/home.png'), size=(20, 20))
        self.img_cadastrar = ctk.CTkImage(light_image=Image.open('./assets/icons/cadastrar.png'), size=(20, 20))
        self.img_documentos = ctk.CTkImage(light_image=Image.open('./assets/icons/documentos.png'), size=(20, 20))
        self.img_balance = ctk.CTkImage(light_image=Image.open('./assets/icons/balance.png'), size=(20, 20))
        self.img_planilhas = ctk.CTkImage(light_image=Image.open('./assets/icons/planilhas.png'), size=(20, 20))
        self.img_perfilUser = ctk.CTkImage(light_image=Image.open('./assets/icons/user.png'), size=(20, 20))
        
        
        self.window.mainloop()
        
        
    
    
    def Home(self):
                                                     #Janela Inicial
        #img
        self.img2_cadastrar = ctk.CTkImage(light_image=Image.open('./assets/icons/cadastrar.png'), size=(80, 80))
        self.img2_documentos = ctk.CTkImage(light_image=Image.open('./assets/icons/documentos.png'), size=(80, 80))
        self.img2_balance = ctk.CTkImage(light_image=Image.open('./assets/icons/balance.png'), size=(80, 80))
        self.img2_planilhas = ctk.CTkImage(light_image=Image.open('./assets/icons/planilhas.png'), size=(80, 80))
        self.img2_perfilUser = ctk.CTkImage(light_image=Image.open('./assets/icons/user.png'), size=(80, 80))
        self.img2_BhpAjuda = ctk.CTkImage(light_image=Image.open('./assets/icons/lampada.png'), size=(80, 80))    
        
        
        #frames
        self.frame_background = ctk.CTkFrame(self.window, fg_color=Elements.CINZA,width=Elements.BG_WIDTH, height=Elements.BG_HEIGHT)
        self.frame_background.place(x=0, y=0)
        
        
        #buttons
        self.button_CadastrarAutor = ctk.CTkButton(self.frame_background, text='Cadastrar cliente', image=self.img2_cadastrar, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Cadastrar_Cliente,width=350, height=250, border_width=1, text_color=Elements.COLOR_TITLE, border_color=Elements.AZUL_MARINHO)
        self.button_CadastrarAutor.place(x=200, y=160)
        self.button_DocAtualizados = ctk.CTkButton(self.frame_background, text='Doc. Atualizados', image=self.img2_documentos, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Documentos,width=350, height=250, text_color=Elements.COLOR_TITLE, border_width=1)
        self.button_DocAtualizados.place(x=570, y=160)
        self.button_Consultas = ctk.CTkButton(self.frame_background, text='Consultas', image=self.img2_balance, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Documentos,width=350, height=250, text_color=Elements.COLOR_TITLE, border_width=1)
        self.button_Consultas.place(x=940, y=160)
        
        self.button_Planilhas_clientes = ctk.CTkButton(self.frame_background, text='Planilhas de clientes', image=self.img2_planilhas, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Documentos,width=350, height=250, text_color=Elements.COLOR_TITLE, border_width=1)
        self.button_Planilhas_clientes.place(x=200, y=420)
        self.button_Perfil_usuario = ctk.CTkButton(self.frame_background, text='Perfil de úsuario', image=self.img2_perfilUser, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Documentos,width=350, height=250, text_color=Elements.COLOR_TITLE, border_width=1)
        self.button_Perfil_usuario.place(x=570, y=420)
        self.button_BhpAjuda = ctk.CTkButton(self.frame_background, text='BHP Ajuda !', image=self.img2_BhpAjuda, fg_color=Elements.COLOR_WHITE,font=(Elements.Font_DEFAULT, 20, 'bold'),
                                               command=self.Documentos,width=350, height=250, text_color=Elements.COLOR_TITLE, border_width=1)
        self.button_BhpAjuda.place(x=940, y=420)
        
        
        #labels
        self.lbl_Logo_BHP = ctk.CTkLabel(self.frame_background, text='BHP', text_color=Elements.COLOR_DARK, fg_color='transparent',font=(Elements.Font_DEFAULT, 20, 'bold'))
        self.lbl_Logo_BHP.place(x=735, y=715)
        self.lbl_Logo_BHP_System = ctk.CTkLabel(self.frame_background, text='System Lawyers', text_color=Elements.COLOR_SYSTEM_LYS, fg_color='transparent',font=(Elements.Font_DEFAULT, 14, 'bold'), height=15)
        self.lbl_Logo_BHP_System.place(x=700, y=740)



    
    def Documentos(self):
        #frames
        self.frame_background = ctk.CTkFrame(self.window, fg_color=Elements.CINZA,width=Elements.BG_WIDTH, height=Elements.BG_HEIGHT)
        self.frame_background.place(x=0, y=0)
        self.frame_Left = ctk.CTkFrame(self.frame_background, fg_color=Elements.AZUL_MARINHO, width=250, height=Elements.BG_HEIGHT,
                                      corner_radius=1)
        self.frame_Left.place(x=0, y=0)
        self.scrollable_frame = ctk.CTkScrollableFrame(self.frame_background, width=1000, height=450,label_text='Documento')
        self.scrollable_frame.place(x=360, y=100)
        self.texto_saida = scrolledtext.ScrolledText(self.scrollable_frame,wrap=tk.WORD, state="disabled",width=600)
        self.texto_saida.pack()


        #buttons
        self.button_Home = ctk.CTkButton(self.frame_Left, text='   Home', image=self.img_Home, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Home.place(x=-17, y=120)
        self.button_Cadastrar_Left = ctk.CTkButton(self.frame_Left, text='   Cadastrar cliente', image=self.img_cadastrar, fg_color='transparent',
                                                   font=(Elements.Font_DEFAULT, 12, 'bold'), command=self.forget)
        self.button_Cadastrar_Left.place(x=10, y=160)
        self.button_Consultas_processos = ctk.CTkButton(self.frame_Left, text='   Consultas de processos', image=self.img_balance, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Consultas_processos.place(x=10, y=208)
        self.button_Planilhas = ctk.CTkButton(self.frame_Left, text='   Planilhas de clientes', image=self.img_planilhas, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Planilhas.place(x=10, y=250)
        self.button_Perfil_User = ctk.CTkButton(self.frame_Left, text='   Perfil de usúario', image=self.img_perfilUser, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Perfil_User.place(x=10, y=295)
        self.button_Contato = ctk.CTkButton(self.frame_Left, text='     Contato', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Contato.place(x=10, y=600)
        self.button_Ajuda = ctk.CTkButton(self.frame_Left, text=' Ajuda', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Ajuda.place(x=10, y=630)
        self.button_SairDaConta = ctk.CTkButton(self.frame_Left, text='              Sair da conta', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_SairDaConta.place(x=10, y=660)
        
        self.button_Abrir_Word = ctk.CTkButton(self.frame_background, text='Abrir Word',text_color=Elements.CINZA, fg_color=Elements.AZUL_MARINHO,
                                         font=(Elements.Font_DEFAULT, 18, 'bold'), width=300, height=60, corner_radius=25, command=self.ler_documento_word)
        self.button_Abrir_Word.place(x=480, y=660)
        self.button_Abrir_Txt = ctk.CTkButton(self.frame_background, text='Abrir Txt',text_color=Elements.CINZA, fg_color=Elements.AZUL_MARINHO,
                                         font=(Elements.Font_DEFAULT, 18, 'bold'), width=300, height=60, corner_radius=25, command=self.ler_documento_txt)
        self.button_Abrir_Txt.place(x=950, y=660)
        
        
        
        
        
    def Cadastrar_Cliente(self):
        
        #imgs
        
        self.img_Home = ctk.CTkImage(light_image=Image.open('./assets/icons/home.png'), size=(20, 20))
        self.img_cadastrar = ctk.CTkImage(light_image=Image.open('./assets/icons/cadastrar.png'), size=(20, 20))
        self.img_documentos = ctk.CTkImage(light_image=Image.open('./assets/icons/documentos.png'), size=(20, 20))
        self.img_balance = ctk.CTkImage(light_image=Image.open('./assets/icons/balance.png'), size=(20, 20))
        self.img_planilhas = ctk.CTkImage(light_image=Image.open('./assets/icons/planilhas.png'), size=(20, 20))
        self.img_perfilUser = ctk.CTkImage(light_image=Image.open('./assets/icons/user.png'), size=(20, 20))
        
        #Frames
        self.frame_background = ctk.CTkFrame(self.window, fg_color=Elements.CINZA,width=Elements.BG_WIDTH, height=Elements.BG_HEIGHT)
        self.frame_background.place(x=0, y=0)
        self.frame_Left = ctk.CTkFrame(self.frame_background, fg_color=Elements.AZUL_MARINHO, width=250, height=Elements.BG_HEIGHT,
                                       corner_radius=1)
        self.frame_Left.place(x=0, y=0)
        
        #Labels
        self.lbl_Cadastro = ctk.CTkLabel(self.frame_background, text='Cadastrar Autor & Réu', text_color=Elements.COLOR_TITLE, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 30, 'bold'))
        self.lbl_Cadastro.place(x=710, y=60)
        self.lbl_Nome = ctk.CTkLabel(self.frame_background, text='NOME COMPLETO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Nome.place(x=320, y=150)
        #self.lbl_Nacionalidade = ctk.CTkLabel(self.frame_background, text='NACIONALIDADE*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         #font=(Elements.Font_DEFAULT, 12, 'bold'))
        #self.lbl_Nacionalidade.place(x=320, y=210)
        self.lbl_DataDeNascimento = ctk.CTkLabel(self.frame_background, text='DATA DE NASCIMENTO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_DataDeNascimento.place(x=320, y=215)
        self.lbl_Profissao = ctk.CTkLabel(self.frame_background, text='PROFISSÃO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Profissao.place(x=320, y=275)
        self.lbl_Bairro = ctk.CTkLabel(self.frame_background, text='BAIRRO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Bairro.place(x=320, y=335)
        self.lbl_Cep = ctk.CTkLabel(self.frame_background, text='CEP*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Cep.place(x=320, y=395)
        self.lbl_SenhaGov = ctk.CTkLabel(self.frame_background, text='SENHA DO GOV*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_SenhaGov.place(x=320, y=460)
        self.lbl_DividasApagar = ctk.CTkLabel(self.frame_background, text='DÍVIDAS A PAGAR*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_DividasApagar.place(x=320, y=520)
        self.lbl_Rg = ctk.CTkLabel(self.frame_background, text='RG*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Rg.place(x=720, y=150)
        self.lbl_Email = ctk.CTkLabel(self.frame_background, text='E-MAIL*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Email.place(x=720, y=215)
        self.lbl_Rua = ctk.CTkLabel(self.frame_background, text='RUA*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Rua.place(x=720, y=275)
        self.lbl_Cidade = ctk.CTkLabel(self.frame_background, text='CIDADE*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Cidade.place(x=720, y=335)
        self.lbl_BeneficioGov = ctk.CTkLabel(self.frame_background, text='BENEFÍCIO DO GOVERNO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_BeneficioGov.place(x=720, y=395)
        self.lbl_WhatsApp = ctk.CTkLabel(self.frame_background, text='WHATSAPP*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_WhatsApp.place(x=720, y=460)
        self.lbl_MeusAcordos = ctk.CTkLabel(self.frame_background, text='MEUS ACORDOS NO SERASA*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_MeusAcordos.place(x=720, y=520)
        self.lbl_Cpf = ctk.CTkLabel(self.frame_background, text='CPF*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Cpf.place(x=1120, y=150)
        self.lbl_EstadoCivil = ctk.CTkLabel(self.frame_background, text='ESTADO CIVIL*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_EstadoCivil.place(x=1120, y=210)
        self.lbl_Complemento = ctk.CTkLabel(self.frame_background, text='COMPLEMENTO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Complemento.place(x=1120, y=275)
        self.lbl_Estado = ctk.CTkLabel(self.frame_background, text='ESTADO*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_Estado.place(x=1120, y=335)
        self.lbl_SenhaDoSerasa = ctk.CTkLabel(self.frame_background, text='SENHA DO SERASA*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_SenhaDoSerasa.place(x=1120, y=395)
        self.lbl_AnalisesCompletas = ctk.CTkLabel(self.frame_background, text='ANÁLISES COMPLETAS*', text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                         font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_AnalisesCompletas.place(x=1120, y=460)
        self.lbl_option_menu = ctk.CTkLabel(self.frame_background, text='EMPRESA CONTRA*',text_color=Elements.COLOR_LABELS, fg_color='transparent',
                                            font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.lbl_option_menu.place(x=1120, y=522)
        self.lbl_Logo_BHP = ctk.CTkLabel(self.frame_background, text='BHP', text_color=Elements.COLOR_TITLE, fg_color='transparent',font=(Elements.Font_DEFAULT, 20, 'bold'))
        self.lbl_Logo_BHP.place(x=845, y=715)
        self.lbl_Logo_BHP_System = ctk.CTkLabel(self.frame_background, text='System Lawyers', text_color=Elements.COLOR_LIGHT, fg_color='transparent',font=(Elements.Font_DEFAULT, 14, 'bold'), height=15)
        self.lbl_Logo_BHP_System.place(x=810, y=740)
        #self.lbl_Home_left = ctk.CTkLabel(self.frame_Left, text='Home', text_color=Elements.COLOR_LIGHT, fg_color='transparent',font=(Elements.Font_DEFAULT, 14, 'bold'))
        #self.lbl_Home_left.place(x=80, y=150)
        
        
        #Entrys
        self.nome_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.nome_Entry.place(x=320, y=175)
        #self.nacionalidade_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        #self.nacionalidade_Entry.place(x=320, y=234)
        self.dataDeNascimento_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.dataDeNascimento_Entry.place(x=320, y=237)
        self.profissao_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.profissao_Entry.place(x=320, y=300)
        self.bairro_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.bairro_Entry.place(x=320, y=360)
        self.cep_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.cep_Entry.place(x=320, y=420)
        self.senhaDoGov_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.senhaDoGov_Entry.place(x=320, y=483)   
        self.dividasApagar_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.dividasApagar_Entry.place(x=320, y=545)  
        self.rg_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.rg_Entry.place(x=720, y=175)
        self.email_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.email_Entry.place(x=720, y=237)
        self.rua_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.rua_Entry.place(x=720, y=300)
        self.cidade_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.cidade_Entry.place(x=720, y=360)
        self.beneficioGov_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.beneficioGov_Entry.place(x=720, y=420)
        self.whatsApp_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.whatsApp_Entry.place(x=720, y=483)   
        self.meusAcordos_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.meusAcordos_Entry.place(x=720, y=545)
        self.cpf_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.cpf_Entry.place(x=1120, y=175)
        self.estadoCivil_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.estadoCivil_Entry.place(x=1120, y=237)
        self.complemento_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.complemento_Entry.place(x=1120, y=300)
        self.estado_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.estado_Entry.place(x=1120, y=360)
        self.senhaDoSerasa_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.senhaDoSerasa_Entry.place(x=1120, y=420)
        self.analisesCompletas_Entry = ctk.CTkEntry(self.frame_background, width=300, height=30)
        self.analisesCompletas_Entry.place(x=1120, y=483)
        self.empresa_entry = ctk.CTkEntry(self.frame_background, width=300, height=30, placeholder_text='Digite o nome da empresa')
        self.empresa_entry.place(x=1120, y=545)
        
        
        #Buttons
        self.button_Home = ctk.CTkButton(self.frame_Left, text='   Home', image=self.img_Home, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'), command=self.Home)
        self.button_Home.place(x=-17, y=120)
        self.button_Cadastrar_Left = ctk.CTkButton(self.frame_Left, text='   Cadastrar cliente', image=self.img_cadastrar, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Cadastrar_Left.place(x=10, y=160)
        self.button_Documentos = ctk.CTkButton(self.frame_Left, text='    Documentos', image=self.img_documentos, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'),
                                               command=self.Documentos)
        self.button_Documentos.place(x=3, y=208)
        self.button_Consultas_processos = ctk.CTkButton(self.frame_Left, text='   Consultas de processos', image=self.img_balance, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Consultas_processos.place(x=10, y=250)
        self.button_Planilhas = ctk.CTkButton(self.frame_Left, text='   Planilhas de clientes', image=self.img_planilhas, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Planilhas.place(x=10, y=295)
        self.button_Perfil_User = ctk.CTkButton(self.frame_Left, text='   Perfil de usúario', image=self.img_perfilUser, fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Perfil_User.place(x=10, y=340)
        self.button_Contato = ctk.CTkButton(self.frame_Left, text='     Contato', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Contato.place(x=10, y=600)
        self.button_Ajuda = ctk.CTkButton(self.frame_Left, text=' Ajuda', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_Ajuda.place(x=10, y=630)
        self.button_SairDaConta = ctk.CTkButton(self.frame_Left, text='              Sair da conta', fg_color='transparent',font=(Elements.Font_DEFAULT, 12, 'bold'))
        self.button_SairDaConta.place(x=10, y=660)
        self.button_Selec_Salvar = ctk.CTkButton(self.frame_background, text='Selecionar e Salvar Doc.',text_color=Elements.CINZA, fg_color=Elements.AZUL_MARINHO,
                                         font=(Elements.Font_DEFAULT, 18, 'bold'), width=300, height=60, corner_radius=25, command=self.selecionar_word_e_escolher_pasta)
        self.button_Selec_Salvar.place(x=320, y=640)
        self.button_Salvar_Bloco_Notas = ctk.CTkButton(self.frame_background, text='Salvar em bloco de notas',text_color=Elements.CINZA, fg_color=Elements.AZUL_MARINHO,
                                         font=(Elements.Font_DEFAULT, 18, 'bold'), width=300, height=60, corner_radius=25, command=self.salvar_como_bloco_de_notas)
        self.button_Salvar_Bloco_Notas.place(x=720, y=640)
        self.button_Limpar = ctk.CTkButton(self.frame_background, text='LIMPAR CAMPOS',text_color=Elements.CINZA, fg_color=Elements.COLOR_DARK,
                                         font=(Elements.Font_DEFAULT, 16, 'bold'), width=300, height=60, corner_radius=25, command=self.limpar_campos, hover_color='#535455')
        self.button_Limpar.place(x=1120, y=640)
        self.sugestoes = tk.Listbox(self.frame_background, width=49,height=5)
        self.sugestoes.place_forget()
        
        self.sugestoes.bind("<ButtonRelease-1>", self.selecionar_empresa)
        self.empresa_entry.bind("<KeyRelease>", self.pesquisar_empresa)
        self.empresa_entry.bind("<FocusIn>", self.pesquisar_empresa)
        
    
    
    def Login(self):
        #img
        self.img_Login = ctk.CTkImage(light_image=Image.open('./assets/icons/direito.png'), size=(750, 770))
        
        #frame
        self.frame_background = ctk.CTkFrame(self.window, fg_color=Elements.COLOR_WHITE,width=Elements.BG_WIDTH, height=Elements.BG_HEIGHT)
        self.frame_background.place(x=0, y=0)
        self.frame_Left = ctk.CTkFrame(self.frame_background, fg_color=Elements.AZUL_MARINHO, width=750, height=Elements.BG_HEIGHT,
                                       corner_radius=1)
        self.frame_Left.place(x=0, y=0)
        
        #label
        self.img_lbl_login = ctk.CTkLabel(self.frame_Left, text='\nBHP \nSystem Lawayers', image=self.img_Login, text_color=Elements.COLOR_WHITE,
                                          font=(Elements.Font_DEFAULT, 60, 'bold'))
        self.img_lbl_login.place(x=0, y=0)
        self.title_Login = ctk.CTkLabel(self.frame_background, text='FAÇA LOGIN', font=(Elements.Font_DEFAULT, 30, 'bold'),
                                        text_color=Elements.COLOR_TITLE)
        self.title_Login.place(x=1050, y=130)
        self.lbl_email = ctk.CTkLabel(self.frame_background, text='DIGITE SEU E-MAIL', font=(Elements.Font_DEFAULT, 13, 'bold'),
                                        text_color=Elements.COLOR_TITLE)
        self.lbl_email.place(x=950, y=255)
        self.lbl_Senha = ctk.CTkLabel(self.frame_background, text='DIGITE SUA SENHA', font=(Elements.Font_DEFAULT, 13, 'bold'),
                                        text_color=Elements.COLOR_TITLE)
        self.lbl_Senha.place(x=950, y=365)
        
        #Entrys
        self.usuario_Entry = ctk.CTkEntry(self.frame_background, width=400, height=40)
        self.usuario_Entry.place(x=950, y=280) 
        self.senha_Entry = ctk.CTkEntry(self.frame_background, width=400, height=40, show='*')
        self.senha_Entry.place(x=950, y=390)
        
        
        #buttons
        self.button_Entrar = ctk.CTkButton(self.frame_background, text='ENTRAR', fg_color=Elements.AZUL_MARINHO, font=(Elements.Font_DEFAULT, 18, 'bold'),
                                           width=300, height=60, corner_radius=50, command=self.Acesso_Login)
        self.button_Entrar.place(x=1010, y=550)
    
    
    
    
    
    
    def pesquisar_empresa(self,event=None):
    # Obter o texto digitado pelo usuário
        texto_digitado = self.empresa_entry.get().lower()

    # Limpar a lista de sugestões
        self.sugestoes.delete(0, 'end')

    # Mostrar sugestões com base no texto digitado
        for empresa in empresas:
            if texto_digitado in empresa.lower():
                self.sugestoes.insert('end', empresa)

    # Exibir a lista de sugestões apenas se houver texto digitado
        if texto_digitado:
            self.sugestoes.place(x=1120, y=573)
        else:
            self.sugestoes.place_forget()
    
    
    
    
    def selecionar_empresa(self,event):
        # Obter o índice da empresa selecionada na lista de sugestões
        indice_selecionado = self.sugestoes.curselection()

        # Verificar se um item foi selecionado na lista
        if indice_selecionado:
            # Obter o nome da empresa selecionada
            empresa_selecionada = self.sugestoes.get(indice_selecionado)
            
            # Preencher o Entry com o nome da empresa selecionada
            self.empresa_entry.delete(0, 'end')
            self.empresa_entry.insert(0, empresa_selecionada)
       
            # Esconder a lista de sugestões
            self.sugestoes.place_forget()
    
    
    
    
    def selecionar_word_e_escolher_pasta(self):
        global company_name
        selected_data_file = "./assets/Empresas.docx"
        
        nome = self.nome_Entry.get()
        nome_maiusculo = nome.upper()
        data_nascimento = self.dataDeNascimento_Entry.get()
        estado_civil = self.estadoCivil_Entry.get()
        profissao = self.profissao_Entry.get()
        rg = self.rg_Entry.get()
        cpf = self.cpf_Entry.get()
        rua = self.rua_Entry.get()
        complemento = self.complemento_Entry.get()
        bairro = self.bairro_Entry.get()
        cidade = self.cidade_Entry.get()
        estado = self.estado_Entry.get()
        numero_cep = self.cep_Entry.get()
        email = self.email_Entry.get()
        
            
        if nome == '' or data_nascimento == '':
            CTkMessagebox(title='BHP System Lawyers',message='Por favor preencha os campos que são obrigatorios !',icon='cancel')
        
        else:
            arquivo = filedialog.askopenfilename(initialdir="/", title="Selecione um arquivo",
                                        filetypes=(("Documentos Word", ".docx"), ("Todos os arquivos", ".*")))
            if arquivo:
                choice =  self.empresa_entry.get()
                data_document = Document(selected_data_file)
                for paragraph in data_document.paragraphs:
                    if choice in paragraph.text:
                        company_name = paragraph.text
                        break
                else:
                    CTkMessagebox(title='BHP System Lawyers',message=f'Não encontrados no arquivo de dados {choice}',icon='cancel')
                    return
        
            model_document = Document(arquivo)
            marker_mapping = {
                r"%Empresa%": company_name,
                r"%nome%": nome_maiusculo,
                r"%estado_civil%": estado_civil,
                r"%cpf%": cpf,
                r"%rg%": rg,
                r"%profissao%": profissao,
                r"%complemento%": complemento,
                r"%bairro%": bairro,
                r"%cidade%": cidade,
                r"%estado%": estado,
                r"%numero_cep%": numero_cep,
                r"%email%": email,
                r"%rua%": rua,
                r"%data_nascimento%": data_nascimento
                
            }
            
            for paragraph in model_document.paragraphs:
                for marker, value in marker_mapping.items():
                    if marker in paragraph.text:
                        paragraph.text = paragraph.text.replace(marker, value)
                        for run in paragraph.runs:
                            run.font.size = Pt(14)  # Define o tamanho da fonte como 14 pontos
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                model_document.save(save_path)
                CTkMessagebox(title='BHP System Lawyers',message=f'Documento Salvo em {save_path}',icon='check')
               
    

    def salvar_como_bloco_de_notas(self):
        destino = filedialog.asksaveasfilename(initialdir="/", title="Salvar arquivo como",
                                          filetypes=(("Arquivos de Texto", ".txt"), ("Todos os arquivos", ".*")),
                                          defaultextension=".txt")
        if destino:    
            nome = self.nome_Entry.get()
            data_nascimento = self.dataDeNascimento_Entry.get()
            estado_civil = self.estadoCivil_Entry.get()
            profissao = self.profissao_Entry.get()
            rg = self.rg_Entry.get()
            cpf = self.cpf_Entry.get()
            beneficio_gov = self.beneficioGov_Entry.get()
            senha_serasa = self.senhaDoSerasa_Entry.get()
            senha_gov = self.senhaDoGov_Entry.get()
            dividas_a_pagar = self.dividasApagar_Entry.get()
            acordos_serasa = self.meusAcordos_Entry.get()
            analises_completas = self.analisesCompletas_Entry.get()
            whatsapp = self.whatsApp_Entry.get()
        
            informacoes_formatadas = f'Nome: {nome}\nEstado Civil: {estado_civil}\nProfissão: {profissao}\nRG: {rg}\nCPF: {cpf}\nData de Nascimento: {data_nascimento}\nWhatsApp: {whatsapp}\n\n---------------------------------------------------------------------\nBeneficio_gov: {beneficio_gov}\nSenha_Serasa: {senha_serasa}\nSenha_gov: {senha_gov}\n\n\nDividas: {dividas_a_pagar}\nAcordos: {acordos_serasa}\n\n---------------------------------------------------------------------\nAnalises Completas: {analises_completas}\n'
            informacoes_formatadas = informacoes_formatadas.upper()
            
            with open(destino, 'w', encoding="utf-8") as arquivo_destino:
                arquivo_destino.write(informacoes_formatadas)
                
                
                CTkMessagebox(title='BHP System Lawyers', message=f'Arquivo salvo em: {destino}',icon='check', option_1='Ok')     

    
    
    def ler_documento_word(self):
        caminho_arquivo = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if caminho_arquivo:
            document = Document(caminho_arquivo)
            conteudo = ""
            for paragrafo in document.paragraphs:
                conteudo += paragrafo.text + "\n"
            self.mostrar_conteudo(conteudo=conteudo)
    
    
    def ler_documento_txt(self):
        arquivo = filedialog.askopenfilename(filetypes=[("Arquivos de Texto", "*.txt")])
        if arquivo:
            with open(arquivo, "r", encoding="utf-8") as file:
                conteudo = file.read()
            self.mostrar_conteudo(conteudo=conteudo)

    
    def mostrar_conteudo(self,conteudo):
        self.texto_saida.config(state="normal") # Habilita a edição
        self.texto_saida.delete("1.0", "end")
        self.texto_saida.insert("end", conteudo)
        self.texto_saida.config(state="disabled") # Desabilita a edição

    def forget(self):
        self.frame_background.place_forget()
        
        
        
    def limpar_campos(self):
        
        msg = CTkMessagebox(title="Exit?", message="Tem certeza que deseja limpar os campos preenchidos ?",
                        icon="question", option_1="Cancel", option_2="Não", option_3="Sim")
        
        response = msg.get()
        
        if response == 'Sim':
            self.nome_Entry.delete(0, "end")
            self.dataDeNascimento_Entry.delete(0, "end")
            self.estadoCivil_Entry.delete(0, "end")
            self.profissao_Entry.delete(0, "end")
            self.rg_Entry.delete(0, "end")
            self.cpf_Entry.delete(0, "end")
            self.rua_Entry.delete(0, "end")
            self.complemento_Entry.delete(0, "end")
            self.bairro_Entry.delete(0, "end")
            self.cidade_Entry.delete(0, "end")
            self.estado_Entry.delete(0, "end")
            self.cep_Entry.delete(0, "end")
            self.email_Entry.delete(0, "end")
            self.beneficioGov_Entry.delete(0, "end")
            self.senhaDoSerasa_Entry.delete(0, "end")
            self.senhaDoGov_Entry.delete(0, "end")
            self.dividasApagar_Entry.delete(0, "end")
            self.meusAcordos_Entry.delete(0, "end")
            self.analisesCompletas_Entry.delete(0, "end")
            self.whatsApp_Entry.delete(0, "end")
            
        else:
            pass
        
   
    def Acesso_Login(self):
        try:
            usuario = self.usuario_Entry.get()
            senha = self.senha_Entry.get()
            mydb = Connect()
            cursor = mydb.cursor()
            sql = 'SELECT * FROM python.tb_usuarios WHERE nome = %s AND senha = %s'
            values = (usuario, senha)
            cursor.execute(sql, values)
            result = cursor.fetchone()
            
            if result:
                CTkMessagebox(title='Sucess', message=f'Bem vindo ! {usuario}', icon='check')
                self.Home()
            
            else:
                CTkMessagebox(title='Error', message='E-mail ou senha incorretos.', icon='cancel')
        
        except Exception as error:
            CTkMessagebox(title='Error', message=f'Algo deu Errado, {error}', icon='cancel')
      
        


if __name__ == "__main__":
    root = App()